from __future__ import annotations

from pathlib import Path

from docx import Document

from app.builders.docx_builder import DocxBuilder
from app.config import load_config
from app.parsers.ppt_parser import PPTParser
from tests.sample_factory import build_sample_pptx


def test_docx_builder_writes_paragraphs_tables_and_images(tmp_path):
    pptx_path = build_sample_pptx(tmp_path / "demo_course.pptx")
    config = load_config()
    parser = PPTParser(config, asset_dir=tmp_path / "assets")
    document_ir = parser.parse(pptx_path)

    output_path = tmp_path / "demo_course.docx"
    DocxBuilder(config).build(document_ir, output_path)

    exported = Document(output_path)
    paragraph_text = "\n".join(paragraph.text for paragraph in exported.paragraphs)

    assert "概率论导论" in paragraph_text
    assert "离散型随机变量" in paragraph_text
    assert exported.tables[0].cell(1, 2).text == "0.4"
    assert len(exported.inline_shapes) >= 2
    assert Path(
        next(
            block for block in document_ir.slides[1].blocks if block.block_type == "diagram"
        ).content["preview_path"]
    ).exists()
