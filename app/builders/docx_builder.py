from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from app.config import AppConfig
from app.models.ir_models import ContentBlock, SlideIRDocument
from app.utils.fonts import apply_run_fonts


class DocxBuilder:
    def __init__(self, config: AppConfig):
        self.config = config

    def build(self, document_ir: SlideIRDocument, output_path: Path) -> Path:
        document = Document()
        self._configure_page(document)
        rendered_captions: set[str] = set()
        for slide_idx, slide in enumerate(document_ir.slides, start=1):
            if slide_idx > 1:
                document.add_page_break()
            if slide.title:
                paragraph = document.add_paragraph(style="Heading 1")
                self._add_plain_run(
                    paragraph, slide.title, size=max(self.config.default_font_size_pt + 4, 16)
                )
            for block in slide.blocks:
                if block.semantic_role == "caption" and block.block_id in rendered_captions:
                    continue
                self._render_block(document, block)
                if block.block_type in {"image", "table", "diagram"}:
                    for relation in block.relations:
                        if relation.relation_type != "has-caption":
                            continue
                        caption = next(
                            (
                                candidate
                                for candidate in slide.blocks
                                if candidate.block_id == relation.target_block_id
                            ),
                            None,
                        )
                        if caption:
                            self._render_caption(document, caption)
                            rendered_captions.add(caption.block_id)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(output_path))
        return output_path

    def _configure_page(self, document: Document) -> None:
        section = document.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        normal_style = document.styles["Normal"]
        normal_style.font.size = Pt(self.config.default_font_size_pt)
        normal_style.font.name = self.config.default_body_font

    def _render_block(self, document: Document, block: ContentBlock) -> None:
        if block.semantic_role == "footer" and not self.config.export_footer:
            return
        if block.block_type == "text":
            self._render_text_block(document, block)
        elif block.block_type == "table":
            self._render_table(document, block)
        elif block.block_type == "image":
            self._render_image(document, block)
        elif block.block_type == "diagram":
            self._render_diagram(document, block)

    def _render_text_block(self, document: Document, block: ContentBlock) -> None:
        paragraphs = block.content.get("paragraphs") or []
        if not paragraphs:
            paragraph = document.add_paragraph()
            self._add_plain_run(paragraph, block.content.get("text", ""))
            return
        for entry in paragraphs:
            style = "Heading 2" if block.semantic_role == "subtitle" else None
            paragraph = document.add_paragraph(style=style)
            paragraph.paragraph_format.line_spacing = self.config.default_line_spacing
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            if entry.get("alignment") == "center":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if entry.get("bullet"):
                paragraph.style = document.styles["List Bullet"]
            for run_data in entry.get("runs", []):
                self._add_styled_run(paragraph, run_data)

    def _render_table(self, document: Document, block: ContentBlock) -> None:
        rows = int(block.content.get("rows", 0))
        cols = int(block.content.get("cols", 0))
        table = document.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        merge_targets: list[tuple[_Cell, _Cell]] = []
        for cell_data in block.content.get("cells", []):
            cell = table.cell(cell_data["row"], cell_data["col"])
            cell.text = cell_data["text"]
            if cell_data["row_span"] > 1 or cell_data["col_span"] > 1:
                target = table.cell(
                    cell_data["row"] + cell_data["row_span"] - 1,
                    cell_data["col"] + cell_data["col_span"] - 1,
                )
                merge_targets.append((cell, target))
        for origin, target in merge_targets:
            origin.merge(target)

    def _render_image(self, document: Document, block: ContentBlock) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(block.content["path"], width=Inches(6.2))

    def _render_diagram(self, document: Document, block: ContentBlock) -> None:
        preview_path = block.content.get("preview_path")
        if preview_path and Path(preview_path).exists():
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(str(preview_path), width=Inches(5.8))
        paragraph = document.add_paragraph()
        text = block.content.get("text") or block.content.get("description") or "Diagram block"
        self._add_plain_run(paragraph, f"[图形降级] {text}")
        if block.content.get("svg_path"):
            meta = document.add_paragraph()
            self._add_plain_run(meta, f"SVG asset: {block.content['svg_path']}", size=10)

    def _render_caption(self, document: Document, block: ContentBlock) -> None:
        paragraph = document.add_paragraph(style="Caption")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_plain_run(paragraph, block.content.get("text", ""))

    def _add_plain_run(self, paragraph: Paragraph, text: str, size: float | None = None) -> Run:
        run = paragraph.add_run(text)
        apply_run_fonts(run, self.config.default_body_font, self.config.default_cjk_font)
        run.font.size = Pt(size or self.config.default_font_size_pt)
        return run

    def _add_styled_run(self, paragraph: Paragraph, run_data: dict) -> None:
        run = paragraph.add_run(run_data.get("text", ""))
        run.font.bold = run_data.get("bold", False)
        run.font.italic = run_data.get("italic", False)
        run.font.underline = run_data.get("underline", False)
        if color := run_data.get("color"):
            run.font.color.rgb = RGBColor.from_string(color)
        run.font.size = Pt(run_data.get("font_size_pt") or self.config.default_font_size_pt)
        apply_run_fonts(
            run,
            run_data.get("font_name") or self.config.default_body_font,
            self.config.default_cjk_font,
        )
